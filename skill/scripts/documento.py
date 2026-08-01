"""Borde de salida: convierte un ResultadoCalculo en .docx (y PDF si hay LibreOffice).

No calcula nada: recibe números ya cerrados y los pinta sobre assets/plantilla.docx.
Los marcadores {{ASI}} de la plantilla están cada uno en un run propio, por eso
el reemplazo por run es suficiente y no hay que reconstruir párrafos.
"""

import shutil
import subprocess
from decimal import Decimal
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor

PLANTILLA = Path(__file__).resolve().parent.parent / "assets" / "plantilla.docx"


def _cop(valor: Decimal) -> str:
    """$ 1.938.000 — separador de miles con punto, sin decimales."""
    return "$ " + f"{int(valor):,}".replace(",", ".")


def _pct(valor: Decimal) -> str:
    return f"{valor * 100:.0f}%"


def _reemplazar_en(parrafos, valores: dict) -> None:
    for p in parrafos:
        for run in p.runs:
            if "{{" in run.text:
                for marcador, texto in valores.items():
                    run.text = run.text.replace(marcador, texto)


def _parrafos_del_documento(doc):
    yield from doc.paragraphs
    for tabla in doc.tables:
        for fila in tabla.rows:
            for celda in fila.cells:
                yield from celda.paragraphs
    for seccion in doc.sections:
        yield from seccion.footer.paragraphs


def _tabla_por_encabezado(doc, primer_encabezado: str):
    for tabla in doc.tables:
        if tabla.rows[0].cells[0].text.strip() == primer_encabezado:
            return tabla
    raise ValueError(f"La plantilla no tiene tabla con encabezado {primer_encabezado!r}")


def _celda(fila, indice, texto, alineacion, negrita=False):
    celda = fila.cells[indice]
    p = celda.paragraphs[0]
    p.alignment = alineacion
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run(texto)
    r.font.size = Pt(9.5)
    r.bold = negrita
    return celda


def _insertar_banda_borrador(doc) -> None:
    """Banda visible al inicio del documento, antes del encabezado."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("BORRADOR — REQUIERE APROBACIÓN DE GERENCIA COMERCIAL")
    r.bold = True
    r.font.size = Pt(12)
    r.font.color.rgb = RGBColor(0x84, 0x20, 0x29)
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:fill"), "F8D7DA")
    p._p.get_or_add_pPr().append(shd)
    # El cuerpo arranca con la tabla del encabezado: mover la banda antes de ella.
    doc.tables[0]._tbl.addprevious(p._p)


def generar_docx(resultado, cliente, numero: str, textos: dict, destino: Path) -> Path:
    """Renderiza la cotización. Devuelve la ruta del .docx.

    Args:
        resultado: ResultadoCalculo ya cerrado (calculo.calcular).
        cliente: dict de datos_io.cargar_clientes.
        numero: consecutivo tipo COT-2026-001 (historial.siguiente_numero).
        textos: {"fecha", "alcance", "tiempo_entrega", "observaciones",
                 "asesor", "validez_dias"} — redacción que aporta el flujo.
        destino: carpeta de salida (fuera de skill/).
    """
    doc = Document(PLANTILLA)

    tiene_retefuente = resultado.retefuente_informativa > 0
    if tiene_retefuente:
        neto = resultado.total - resultado.retefuente_informativa
        bloque_rete = (
            "Nota informativa: el cliente, como agente retenedor, practicará "
            f"retención en la fuente por compras (2,5%) de {_cop(resultado.retefuente_informativa)} "
            f"al momento del pago. Valor neto a recibir estimado: {_cop(neto)}. "
            "La retención no se descuenta del total de esta cotización."
        )
    else:
        bloque_rete = ""

    valores = {
        "{{NUMERO}}": numero,
        "{{FECHA}}": textos.get("fecha", ""),
        "{{CLIENTE_NOMBRE}}": cliente.get("razon_social", ""),
        "{{CLIENTE_NIT}}": cliente.get("nit", ""),
        "{{CLIENTE_CONTACTO}}": cliente.get("contacto", ""),
        "{{CLIENTE_CIUDAD}}": cliente.get("ciudad", ""),
        "{{CONDICION_PAGO}}": cliente.get("condicion_pago", ""),
        "{{VALIDEZ_DIAS}}": str(textos.get("validez_dias", "")),
        "{{ALCANCE}}": textos.get("alcance", ""),
        "{{TIEMPO_ENTREGA}}": textos.get("tiempo_entrega", ""),
        "{{OBSERVACIONES}}": textos.get("observaciones", "Ninguna."),
        "{{ASESOR}}": textos.get("asesor", "Equipo comercial"),
        "{{SUBTOTAL}}": _cop(resultado.subtotal),
        "{{DESCUENTO_MANUAL_PCT}}": _pct(resultado.descuento_pct),
        "{{DESCUENTO_MANUAL}}": _cop(resultado.descuento_valor),
        "{{IVA}}": _cop(resultado.iva),
        "{{TOTAL}}": _cop(resultado.total),
        "{{RETEFUENTE_BLOQUE}}": bloque_rete,
    }
    _reemplazar_en(_parrafos_del_documento(doc), valores)

    items = _tabla_por_encabezado(doc, "Código")
    for linea in resultado.lineas:
        fila = items.add_row()
        _celda(fila, 0, linea.sku, WD_ALIGN_PARAGRAPH.CENTER)
        _celda(fila, 1, linea.nombre, WD_ALIGN_PARAGRAPH.LEFT)
        _celda(fila, 2, str(linea.cantidad), WD_ALIGN_PARAGRAPH.CENTER)
        _celda(fila, 3, _cop(linea.precio_unitario), WD_ALIGN_PARAGRAPH.RIGHT)
        pct = _pct(linea.descuento_volumen_pct) if linea.descuento_volumen_pct else "—"
        _celda(fila, 4, pct, WD_ALIGN_PARAGRAPH.CENTER)
        _celda(fila, 5, _cop(linea.subtotal_con_descuento), WD_ALIGN_PARAGRAPH.RIGHT)

    if resultado.requiere_aprobacion:
        _insertar_banda_borrador(doc)

    destino.mkdir(parents=True, exist_ok=True)
    ruta = destino / f"{numero}.docx"
    doc.save(ruta)
    return ruta


def a_pdf(ruta_docx: Path) -> Path:
    """Convierte con LibreOffice si está disponible; si no, devuelve el .docx."""
    soffice = shutil.which("soffice") or shutil.which("libreoffice")
    if soffice is None:
        return ruta_docx
    try:
        subprocess.run(
            [soffice, "--headless", "--convert-to", "pdf",
             "--outdir", str(ruta_docx.parent), str(ruta_docx)],
            capture_output=True, timeout=120, check=True,
        )
    except (subprocess.SubprocessError, OSError):
        return ruta_docx
    ruta_pdf = ruta_docx.with_suffix(".pdf")
    return ruta_pdf if ruta_pdf.exists() else ruta_docx
