"""Tests del renderizador: generan .docx reales en tmp_path y los releen."""

import sys
from decimal import Decimal
from pathlib import Path

from docx import Document

sys.path.insert(0, str(Path(__file__).resolve().parent.parent / "skill" / "scripts"))

import calculo  # noqa: E402
import documento  # noqa: E402

CATALOGO = {
    "SIL-006": {"nombre": "Silla rimax apilable", "precio_unitario": Decimal("68000"), "stock": 240},
    "ESC-001": {"nombre": "Escritorio Cedro 120x60", "precio_unitario": Decimal("489000"), "stock": 40},
}
POLITICAS = {
    "descuentos_volumen": [
        {"desde": 1, "hasta": 19, "descuento": Decimal("0.00")},
        {"desde": 20, "hasta": 49, "descuento": Decimal("0.05")},
        {"desde": 50, "hasta": None, "descuento": Decimal("0.10")},
    ],
    "descuento_manual_maximo": Decimal("0.15"),
    "validez_dias": 15,
}
CLIENTE = {
    "nit": "830214569-7", "razon_social": "Fiduciaria Andina SA",
    "contacto": "Comité de Compras", "ciudad": "Bogotá",
    "agente_retenedor": True, "condicion_pago": "30 dias", "notas": "",
}
TEXTOS = {
    "fecha": "2026-07-30", "alcance": "Dotación piso nuevo.",
    "tiempo_entrega": "2 días hábiles (Bogotá)", "observaciones": "Sin novedades.",
    "asesor": "Vendedor Demo", "validez_dias": 15,
}


def _texto_completo(ruta):
    doc = Document(ruta)
    partes = [p.text for p in doc.paragraphs]
    for t in doc.tables:
        for fila in t.rows:
            partes.extend(c.text for c in fila.cells)
    for s in doc.sections:
        partes.extend(p.text for p in s.footer.paragraphs)
    return "\n".join(partes)


def _generar(tmp_path, items, cliente=CLIENTE, descuento=None):
    r = calculo.calcular(items, CATALOGO, cliente, POLITICAS, descuento_manual_pct=descuento)
    return documento.generar_docx(r, cliente, "COT-2026-001", TEXTOS, tmp_path), r


def test_genera_docx_sin_marcadores_pendientes(tmp_path):
    ruta, _ = _generar(tmp_path, [{"sku": "SIL-006", "cantidad": 30}])
    assert ruta == tmp_path / "COT-2026-001.docx"
    texto = _texto_completo(ruta)
    assert "{{" not in texto  # todos los marcadores fueron reemplazados
    assert "Fiduciaria Andina SA" in texto
    assert "COT-2026-001" in texto


def test_lineas_y_totales_en_el_documento(tmp_path):
    # 30 × 68.000 = 2.040.000, 5% → 1.938.000; IVA 368.220; total 2.306.220
    ruta, _ = _generar(tmp_path, [{"sku": "SIL-006", "cantidad": 30}])
    texto = _texto_completo(ruta)
    assert "SIL-006" in texto and "5%" in texto
    assert "$ 1.938.000" in texto
    assert "$ 368.220" in texto
    assert "$ 2.306.220" in texto


def test_retefuente_informativa_cuando_aplica(tmp_path):
    # base 11.613.750 > 1.000.000 y retenedor → nota con 290.344; total intacto.
    ruta, r = _generar(tmp_path, [{"sku": "ESC-001", "cantidad": 25}])
    texto = _texto_completo(ruta)
    assert "$ 290.344" in texto
    assert "no se descuenta del total" in texto
    assert documento._cop(r.total) in texto


def test_sin_retefuente_no_queda_nota(tmp_path):
    cliente = dict(CLIENTE, agente_retenedor=False)
    ruta, _ = _generar(tmp_path, [{"sku": "SIL-006", "cantidad": 10}], cliente=cliente)
    assert "retención" not in _texto_completo(ruta)


def test_banda_borrador_solo_si_requiere_aprobacion(tmp_path):
    banda = "BORRADOR — REQUIERE APROBACIÓN"
    con, _ = _generar(tmp_path, [{"sku": "SIL-006", "cantidad": 30}], descuento=Decimal("0.20"))
    assert banda in _texto_completo(con)
    # La banda queda de primera en el cuerpo del documento.
    assert banda in Document(con).paragraphs[0].text
    (tmp_path / "sin").mkdir()
    sin = documento.generar_docx(
        calculo.calcular([{"sku": "SIL-006", "cantidad": 30}], CATALOGO, CLIENTE, POLITICAS),
        CLIENTE, "COT-2026-002", TEXTOS, tmp_path / "sin",
    )
    assert banda not in _texto_completo(sin)


def test_a_pdf_sin_libreoffice_devuelve_docx(tmp_path, monkeypatch):
    monkeypatch.setattr(documento.shutil, "which", lambda _: None)
    ruta, _ = _generar(tmp_path, [{"sku": "SIL-006", "cantidad": 1}])
    assert documento.a_pdf(ruta) == ruta  # no falla, degrada al .docx


# --- ficha técnica de producto ---------------------------------------------

PRODUCTO = {
    "nombre": "Estante metálico carga liviana 200x90x40",
    "categoria": "archivadores", "unidad": "unidad",
    "precio_unitario": Decimal("345000"), "stock": 52,
}
ESPECIFICACIONES = {
    "dimensiones": "200 × 90 × 40 cm", "materiales": ["Metal"], "caracteristicas": [],
}
TEXTOS_FICHA = {
    "fecha": "2026-08-01", "tiempo_entrega": "2 días hábiles (Bogotá)",
    "observaciones": "Ninguna.", "asesor": "Vendedor Demo", "validez_dias": 15,
    "nota_especificaciones": "Especificaciones tomadas del catálogo vigente.",
}


def test_ficha_se_nombra_por_sku_y_no_deja_marcadores(tmp_path):
    ruta = documento.generar_ficha_docx(
        "ARC-007", PRODUCTO, ESPECIFICACIONES, TEXTOS_FICHA, tmp_path)
    assert ruta == tmp_path / "FICHA-ARC-007.docx"  # sin consecutivo COT-
    texto = _texto_completo(ruta)
    assert "{{" not in texto
    assert "FICHA TÉCNICA" in texto and "ARC-007" in texto


def test_ficha_muestra_specs_precio_y_disponibilidad(tmp_path):
    ruta = documento.generar_ficha_docx(
        "ARC-007", PRODUCTO, ESPECIFICACIONES, TEXTOS_FICHA, tmp_path)
    texto = _texto_completo(ruta)
    assert "200 × 90 × 40 cm" in texto
    assert "Metal" in texto
    assert "$ 345.000" in texto  # precio de lista antes de IVA
    assert "Disponible" in texto and "52 unidades" in texto
    assert "2 días hábiles (Bogotá)" in texto


def test_ficha_declara_lo_que_el_catalogo_no_registra(tmp_path):
    vacias = {"dimensiones": "", "materiales": [], "caracteristicas": []}
    textos = dict(TEXTOS_FICHA,
                  nota_especificaciones="El catálogo no registra dimensiones, "
                                        "materiales y características para este producto.")
    ruta = documento.generar_ficha_docx(
        "SIL-009", dict(PRODUCTO, nombre="Butaco giratorio para laboratorio"),
        vacias, textos, tmp_path)
    texto = _texto_completo(ruta)
    # Concordancia por campo, no un "No registrada" genérico en las tres filas.
    assert "No registradas en catálogo" in texto   # dimensiones y características
    assert "No registrados en catálogo" in texto   # materiales
    assert "El catálogo no registra dimensiones, materiales y características" in texto


def test_ficha_sin_stock_lo_dice(tmp_path):
    ruta = documento.generar_ficha_docx(
        "ARC-008", dict(PRODUCTO, stock=0), ESPECIFICACIONES, TEXTOS_FICHA, tmp_path)
    texto = _texto_completo(ruta)
    assert "Sin existencias" in texto
    assert "Agotado — sujeto a reposición" in texto


def test_ficha_pluraliza_la_unidad_de_venta(tmp_path):
    ruta = documento.generar_ficha_docx(
        "ILU-010", dict(PRODUCTO, unidad="juego", stock=71),
        ESPECIFICACIONES, TEXTOS_FICHA, tmp_path)
    assert "71 juegos" in _texto_completo(ruta)
