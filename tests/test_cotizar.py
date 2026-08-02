"""Tests del orquestador de punta a punta: datos reales, salidas en tmp_path."""

import csv
import json
import sys
from pathlib import Path

import pytest
from docx import Document

sys.path.insert(0, str(Path(__file__).resolve().parent.parent / "skill" / "scripts"))

import cotizar  # noqa: E402
import datos_io  # noqa: E402


def _correr(tmp_path, capsys, peticion, extra_args=()):
    entrada = tmp_path / "peticion.json"
    entrada.write_text(json.dumps(peticion, ensure_ascii=False), encoding="utf-8")
    codigo = cotizar.main(["--entrada", str(entrada), "--salidas", str(tmp_path / "salidas"),
                           *extra_args])
    return codigo, json.loads(capsys.readouterr().out)


PETICION_BASE = {
    "cliente": "café y punto",
    "items": [{"sku": "SIL-006", "cantidad": 30}],
    "fecha": "2026-07-30",
}


def test_flujo_completo_genera_documento_e_historial(tmp_path, capsys):
    codigo, salida = _correr(tmp_path, capsys, PETICION_BASE)
    assert codigo == 0
    assert salida["estado"] == "ok"
    assert salida["numero"] == "COT-2026-001"
    assert salida["cliente"]["nit"] == "901238450-6"
    # 30 × 68.000 × 0,95 = 1.938.000; IVA 368.220; total 2.306.220
    assert salida["totales"]["subtotal"] == "1938000"
    assert salida["totales"]["total"] == "2306220"
    assert Path(salida["documento"]).exists()
    with open(tmp_path / "salidas" / "historial.csv", newline="", encoding="utf-8-sig") as f:
        filas = list(csv.DictReader(f))
    assert filas[0]["numero"] == "COT-2026-001"
    assert filas[0]["total"] == "2306220"


def test_consecutivo_avanza_entre_corridas(tmp_path, capsys):
    _correr(tmp_path, capsys, PETICION_BASE)
    _, segunda = _correr(tmp_path, capsys, PETICION_BASE)
    assert segunda["numero"] == "COT-2026-002"


def test_solo_calculo_no_toca_estado(tmp_path, capsys):
    codigo, salida = _correr(tmp_path, capsys, PETICION_BASE, ["--solo-calculo"])
    assert codigo == 0
    assert salida["estado"] == "ok"
    assert salida["numero"] is None
    assert salida["documento"] is None
    assert not (tmp_path / "salidas").exists()  # ni documento ni historial


def test_sku_no_encontrado_es_error_json(tmp_path, capsys):
    peticion = dict(PETICION_BASE, items=[{"sku": "VEN-999", "cantidad": 1}])
    codigo, salida = _correr(tmp_path, capsys, peticion)
    assert codigo == 1
    assert salida["estado"] == "error"
    assert salida["tipo"] == "sku_no_encontrado"
    assert salida["sku"] == "VEN-999"
    assert not (tmp_path / "salidas").exists()  # no consumió consecutivo


def test_cliente_ambiguo_devuelve_candidatos(tmp_path, capsys):
    # "clínica" coincide con la odontológica y la veterinaria.
    codigo, salida = _correr(tmp_path, capsys, dict(PETICION_BASE, cliente="clínica"))
    assert codigo == 1
    assert salida["tipo"] == "cliente_ambiguo"
    assert len(salida["candidatos"]) >= 2


def test_cliente_no_encontrado(tmp_path, capsys):
    codigo, salida = _correr(tmp_path, capsys,
                             dict(PETICION_BASE, cliente="Astilleros del Pacífico Sur"))
    assert codigo == 1
    assert salida["tipo"] == "cliente_no_encontrado"


# --- cliente nuevo ----------------------------------------------------------
#
# Estos cuatro nombres NO están en clientes.csv. Cada uno se parecía lo
# suficiente a un cliente real como para quedar de candidato único y ser
# adoptado sin preguntar: la cotización salía con el NIT, la ciudad y la
# condición tributaria de otra empresa, numerada y anotada en el historial.

NOMBRES_CLIENTE_NUEVO = (
    "Constructora Sierra Alta SAS",        # se parecía a Constructora Altamira SAS
    "Comercializadora Andina del Norte SAS",  # a Comercializadora Galeras SAS
    "Panadería La Espiga",                 # a Ganadería La Pradera SAS
    "Hotel Playa Dorada",                  # a Hotel Bahía Serena SAS
)


@pytest.mark.parametrize("nombre", NOMBRES_CLIENTE_NUEVO)
def test_cliente_parecido_nunca_se_adopta_en_silencio(tmp_path, capsys, nombre):
    codigo, salida = _correr(tmp_path, capsys, dict(PETICION_BASE, cliente=nombre))
    assert codigo == 1
    assert salida["estado"] == "error"
    assert salida["tipo"] == "cliente_ambiguo"
    # Se listan para que el usuario decida, con el puntaje a la vista.
    assert salida["candidatos"]
    assert all(c["puntaje"] < datos_io.UMBRAL_CERTEZA for c in salida["candidatos"])
    assert not (tmp_path / "salidas").exists()  # ni documento ni consecutivo


CLIENTE_NUEVO = {
    "nit": "901999888-1",
    "razon_social": "Inversiones La Floresta SAS",
    "contacto": "Ana Ruiz",
    "ciudad": "Bogotá",
    "agente_retenedor": False,
}


def _texto(ruta) -> str:
    doc = Document(ruta)
    return "\n".join(
        [p.text for p in doc.paragraphs]
        + [c.text for t in doc.tables for f in t.rows for c in f.cells]
    )


def test_cliente_inline_cotiza_sin_darlo_de_alta(tmp_path, capsys):
    codigo, salida = _correr(tmp_path, capsys, dict(PETICION_BASE, cliente=CLIENTE_NUEVO))
    assert codigo == 0
    assert salida["cliente"] == {"nit": "901999888-1",
                                 "razon_social": "Inversiones La Floresta SAS",
                                 "nuevo": True}
    assert Path(salida["documento"]).exists()
    assert "Inversiones La Floresta SAS" in _texto(salida["documento"])
    # El alta formal la hace cartera: la skill no escribe en su fuente de datos.
    assert "901999888-1" not in datos_io.cargar_clientes()


def test_cliente_inline_se_cotiza_de_contado_aunque_pida_credito(tmp_path, capsys):
    # Políticas §5: primera compra de cliente nuevo siempre de contado.
    peticion = dict(PETICION_BASE, cliente=dict(CLIENTE_NUEVO, condicion_pago="30 dias"))
    codigo, salida = _correr(tmp_path, capsys, peticion)
    assert codigo == 0
    texto = _texto(salida["documento"])
    assert "contado" in texto
    assert "30 dias" not in texto
    assert any("contado" in a and "§5" in a for a in salida["alertas"])


def test_cliente_inline_sin_declarar_retencion_pregunta(tmp_path, capsys):
    # Suponer que no retiene es inventarle una condición tributaria al cliente.
    sin_declarar = {k: v for k, v in CLIENTE_NUEVO.items() if k != "agente_retenedor"}
    codigo, salida = _correr(tmp_path, capsys, dict(PETICION_BASE, cliente=sin_declarar))
    assert codigo == 1
    assert salida["tipo"] == "retencion_no_declarada"
    assert not (tmp_path / "salidas").exists()


def test_cliente_inline_retenedor_declarado_liquida_retefuente(tmp_path, capsys):
    peticion = dict(PETICION_BASE, cliente=dict(CLIENTE_NUEVO, agente_retenedor=True))
    codigo, salida = _correr(tmp_path, capsys, peticion)
    assert codigo == 0
    # Base 1.938.000 > 1.000.000 → 2,5% = 48.450, informativa (no resta del total).
    assert salida["totales"]["retefuente_informativa"] == "48450"
    assert salida["totales"]["total"] == "2306220"


def test_retenedor_en_texto_no_se_lee_al_reves():
    # bool("no") es True: leerlo así le inventaría una retención al cliente.
    assert cotizar._leer_retenedor("no") is False
    assert cotizar._leer_retenedor("sí") is True
    assert cotizar._leer_retenedor(False) is False
    # Lo que no es un sí ni un no claro es un dato ausente, y se pregunta.
    assert cotizar._leer_retenedor("quizás") is None
    assert cotizar._leer_retenedor("") is None
    assert cotizar._leer_retenedor(None) is None


def test_cliente_inline_incompleto_dice_que_falta(tmp_path, capsys):
    incompleto = {"nit": "901999888-1", "razon_social": "Inversiones La Floresta SAS",
                  "agente_retenedor": False}
    codigo, salida = _correr(tmp_path, capsys, dict(PETICION_BASE, cliente=incompleto))
    assert codigo == 1
    assert salida["tipo"] == "cliente_nuevo_incompleto"
    assert salida["campos_faltantes"] == ["contacto", "ciudad"]
    assert not (tmp_path / "salidas").exists()


def test_cliente_inline_con_nit_ya_registrado_no_es_nuevo(tmp_path, capsys):
    # Aceptarlo pisaría la condición de pago y las notas del cliente real.
    suplantador = dict(CLIENTE_NUEVO, nit="901238450-6")
    codigo, salida = _correr(tmp_path, capsys, dict(PETICION_BASE, cliente=suplantador))
    assert codigo == 1
    assert salida["tipo"] == "cliente_ya_registrado"
    assert salida["cliente_registrado"]["razon_social"] == "Café y Punto SAS"
    assert not (tmp_path / "salidas").exists()


def test_stock_insuficiente_alerta_sin_alterar_precios(tmp_path, capsys):
    # SIL-003 (stock 0 en catalogo.csv): 2 × 1.185.000 = 2.370.000 se cotiza pleno.
    peticion = dict(PETICION_BASE, items=[{"sku": "SIL-003", "cantidad": 2}])
    codigo, salida = _correr(tmp_path, capsys, peticion, ["--solo-calculo"])
    assert codigo == 0
    assert salida["totales"]["subtotal"] == "2370000"
    assert any("SIL-003" in a and "reposición" in a for a in salida["alertas"])


def test_descuento_excedido_marca_borrador_en_historial(tmp_path, capsys):
    peticion = dict(PETICION_BASE, descuento_manual_pct="0.20")
    _, salida = _correr(tmp_path, capsys, peticion)
    assert salida["requiere_aprobacion"] is True
    with open(tmp_path / "salidas" / "historial.csv", newline="", encoding="utf-8-sig") as f:
        assert list(csv.DictReader(f))[0]["estado"] == "borrador_requiere_aprobacion"


# --- modo ficha técnica -----------------------------------------------------

PETICION_FICHA = {"skus": ["ARC-007", "SIL-008"], "fecha": "2026-08-01"}


def test_ficha_genera_un_documento_por_sku(tmp_path, capsys):
    codigo, salida = _correr(tmp_path, capsys, PETICION_FICHA, ["--modo", "ficha"])
    assert codigo == 0
    assert salida["estado"] == "ok" and salida["modo"] == "ficha"
    assert [f["sku"] for f in salida["fichas"]] == ["ARC-007", "SIL-008"]
    for ficha in salida["fichas"]:
        assert Path(ficha["documento"]).exists()
        assert Path(ficha["documento"]).name.startswith(f"FICHA-{ficha['sku']}")
    assert salida["fichas"][0]["especificaciones"]["dimensiones"] == "200 × 90 × 40 cm"
    assert salida["fichas"][0]["precio_lista"] == "345000"


def test_ficha_no_toca_el_historial(tmp_path, capsys):
    codigo, salida = _correr(tmp_path, capsys, PETICION_FICHA, ["--modo", "ficha"])
    assert codigo == 0
    assert salida["consecutivo_consumido"] is False
    assert not (tmp_path / "salidas" / "historial.csv").exists()


def test_ficha_no_consume_consecutivo_entre_cotizaciones(tmp_path, capsys):
    """La ficha intercalada no debe correr la numeración de cotizaciones."""
    _, primera = _correr(tmp_path, capsys, PETICION_BASE)
    _correr(tmp_path, capsys, PETICION_FICHA, ["--modo", "ficha"])
    _, segunda = _correr(tmp_path, capsys, PETICION_BASE)
    assert primera["numero"] == "COT-2026-001"
    assert segunda["numero"] == "COT-2026-002"  # no saltó a 003
    with open(tmp_path / "salidas" / "historial.csv", newline="", encoding="utf-8-sig") as f:
        filas = list(csv.DictReader(f))
    assert [f_["numero"] for f_ in filas] == ["COT-2026-001", "COT-2026-002"]


def test_ficha_acepta_modo_en_la_peticion_y_sku_singular(tmp_path, capsys):
    peticion = {"modo": "ficha", "sku": "ARC-007", "fecha": "2026-08-01"}
    codigo, salida = _correr(tmp_path, capsys, peticion)
    assert codigo == 0
    assert [f["sku"] for f in salida["fichas"]] == ["ARC-007"]


def test_ficha_alerta_cuando_el_catalogo_no_registra_specs(tmp_path, capsys):
    _, salida = _correr(tmp_path, capsys, PETICION_FICHA, ["--modo", "ficha"])
    alerta = next(a for a in salida["alertas"] if a.startswith("SIL-008"))
    assert "dimensiones y materiales" in alerta
    assert "NO estimar" in alerta


def test_ficha_con_sku_inexistente_pregunta_no_inventa(tmp_path, capsys):
    peticion = {"skus": ["ARC-007", "VEN-999"], "fecha": "2026-08-01"}
    codigo, salida = _correr(tmp_path, capsys, peticion, ["--modo", "ficha"])
    assert codigo == 1
    assert salida["tipo"] == "sku_no_encontrado"
    assert salida["skus"] == ["VEN-999"]
    assert not (tmp_path / "salidas").exists()  # ni siquiera la ficha válida


def _texto_de_tablas(ruta):
    return "\n".join(c.text for t in Document(ruta).tables
                     for f in t.rows for c in f.cells)


def test_ficha_usa_la_ciudad_para_el_tiempo_de_entrega(tmp_path, capsys):
    peticion = {"sku": "ARC-007", "ciudad": "Cúcuta", "fecha": "2026-08-01"}
    _, salida = _correr(tmp_path, capsys, peticion, ["--modo", "ficha"])
    assert "5 días hábiles (Cúcuta)" in _texto_de_tablas(salida["fichas"][0]["documento"])


def test_ficha_sin_ciudad_no_supone_un_destino(tmp_path, capsys):
    _, salida = _correr(tmp_path, capsys, {"sku": "ARC-007"}, ["--modo", "ficha"])
    texto = _texto_de_tablas(salida["fichas"][0]["documento"])
    assert "Según ciudad de destino" in texto
    assert "confirmar con logística" in texto


def test_solo_calculo_no_aplica_a_ficha(tmp_path, capsys):
    codigo, salida = _correr(tmp_path, capsys, PETICION_FICHA,
                             ["--modo", "ficha", "--solo-calculo"])
    assert codigo == 1
    assert salida["tipo"] == "peticion_invalida"


def test_ficha_sin_skus_es_error(tmp_path, capsys):
    codigo, salida = _correr(tmp_path, capsys, {"fecha": "2026-08-01"}, ["--modo", "ficha"])
    assert codigo == 1
    assert salida["tipo"] == "peticion_invalida"
